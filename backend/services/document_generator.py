import os
import markdown2
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageTemplate, Frame
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from ebooklib import epub
from xml.sax.saxutils import escape

def clean_for_pdf(text):
    if not text: return ""
    # Escape XML characters
    text = escape(text)
    # Restore common line break tags if necessary, or ensure they are self-closing
    text = text.replace("&lt;br&gt;", "<br/>").replace("&lt;br/&gt;", "<br/>")
    return text

def generate_documents(results: dict, video_data: dict, output_dir: str, options: list) -> dict:
    """
    Generates the requested files and returns a dictionary of file paths.
    """
    generated_files = {}
    
    # 0. Global Request Log
    log_path = os.path.join(os.path.dirname(output_dir), "request_log.md")
    log_entry = f"| {datetime.now().strftime('%Y-%m-%d %H:%M')} | [{video_data['title']}]({video_data.get('webpage_url', 'Unknown URL')}) |\n"
    
    if not os.path.exists(log_path):
        with open(log_path, "w", encoding="utf-8") as f:
            f.write("# Video Request Log\n\n| Date | Video |\n|---|---|\n")
            
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(log_entry)

    # Base filename
    base_filename = f"{video_data['title'][:30].replace(' ', '_')}_{video_data.get('upload_date', '')}"
    base_filename = "".join([c for c in base_filename if c.isalnum() or c in ('_', '-')])
    
    # Metadata Block for Text/DOCX
    keywords = results.get("keywords", "No keywords generated.")
    metadata_text = (
        f"Title: {video_data['title']}\n"
        f"URL: {video_data.get('webpage_url', '')}\n"
        f"Author: {video_data.get('uploader', 'Unknown')}\n"
        f"Upload Date: {video_data.get('upload_date', '')}\n"
        f"Report Date: {datetime.now().strftime('%Y-%m-%d')}\n"
        f"Keywords: {keywords}\n"
        f"{'-'*40}\n\n"
    )

    # Helper for Doc Generation
    def generate_format(key, content, title_suffix):
        # Markdown
        md_path = os.path.join(output_dir, f"{base_filename}_{key}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {video_data['title']} - {title_suffix}\n\n{metadata_text}{content}")
        generated_files[f"{key}_md"] = f"/download/{os.path.basename(output_dir)}/{os.path.basename(md_path)}"
        
        # DOCX
        docx_path = os.path.join(output_dir, f"{base_filename}_{key}.docx")
        create_docx(f"{video_data['title']} - {title_suffix}", metadata_text + content, docx_path)
        generated_files[f"{key}_docx"] = f"/download/{os.path.basename(output_dir)}/{os.path.basename(docx_path)}"
        
        # PDF (Professional)
        pdf_path = os.path.join(output_dir, f"{base_filename}_{key}.pdf")
        doc_title = f"{title_suffix}: {video_data['title']}"
        create_professional_pdf(doc_title, content, pdf_path, video_data, keywords)
        generated_files[f"{key}_pdf"] = f"/download/{os.path.basename(output_dir)}/{os.path.basename(pdf_path)}"
        
        # EPUB
        epub_path = os.path.join(output_dir, f"{base_filename}_{key}.epub")
        create_epub(f"{title_suffix}: {video_data['title']}", metadata_text + content, epub_path)
        generated_files[f"{key}_epub"] = f"/download/{os.path.basename(output_dir)}/{os.path.basename(epub_path)}"

    if "summary" in results:
        generate_format("summary", results["summary"], "Summary")
        
    if "guide" in results:
        generate_format("guide", results["guide"], "Didactic Guide")

    # Transcriptions (keep simple text)
    if "transcription_orig" in results:
        path = os.path.join(output_dir, f"{base_filename}_transcription_original.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"{metadata_text}{results['transcription_orig']}")
        generated_files["transcription_orig"] = f"/download/{os.path.basename(output_dir)}/{os.path.basename(path)}"

    if "transcription_es" in results:
        path = os.path.join(output_dir, f"{base_filename}_transcription_es.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"{metadata_text}{results['transcription_es']}")
        generated_files["transcription_es"] = f"/download/{os.path.basename(output_dir)}/{os.path.basename(path)}"

    return generated_files

def create_docx(title, content, path):
    doc = Document()
    doc.add_heading(title, 0)
    for paragraph in content.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    doc.save(path)

def create_professional_pdf(doc_name, content, path, video_data, keywords):
    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    
    # Custom Styles
    styles.add(ParagraphStyle(name='Justified', parent=styles['Normal'], alignment=TA_JUSTIFY, spaceAfter=6))
    styles.add(ParagraphStyle(name='Centered', parent=styles['Normal'], alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name='MetaLabel', parent=styles['Normal'], fontSize=8, textColor='gray'))
    styles.add(ParagraphStyle(name='MetaValue', parent=styles['Normal'], fontSize=10, spaceAfter=2))
    
    story = []
    
    # Title Block
    story.append(Paragraph(clean_for_pdf(video_data['title']), styles['Title']))
    story.append(Spacer(1, 12))
    
    # Metadata
    meta = [
        ("URL", video_data.get('webpage_url', '')),
        ("Author", video_data.get('uploader', '')),
        ("Date", video_data.get('upload_date', '')),
        ("Report Date", datetime.now().strftime('%Y-%m-%d')),
        ("Keywords", keywords)
    ]
    
    for label, value in meta:
        story.append(Paragraph(f"<b>{label}:</b> {clean_for_pdf(str(value))}", styles['MetaValue']))
    
    story.append(Spacer(1, 24))
    story.append(Paragraph("_"*50, styles['Centered'])) # Separator
    story.append(Spacer(1, 24))
    
    # Content
    for line in content.split('\n'):
        line = line.strip()
        if line.startswith('### '):
            story.append(Paragraph(clean_for_pdf(line.replace('### ', '')), styles['Heading2']))
        elif line.startswith('## '):
            story.append(Paragraph(clean_for_pdf(line.replace('## ', '')), styles['Heading3']))
        elif line.startswith('# '):
            story.append(Paragraph(clean_for_pdf(line.replace('# ', '')), styles['Heading1']))
        elif line:
            # Handle bullet points simply
            if line.startswith('- '):
                story.append(Paragraph(f"• {clean_for_pdf(line[2:])}", styles['Justified']))
            else:
                story.append(Paragraph(clean_for_pdf(line), styles['Justified']))
    
    # Header/Footer Callback
    def on_page(canvas, doc):
        canvas.saveState()
        
        # Header
        canvas.setFont('Helvetica', 8)
        canvas.drawString(inch, 11*inch, doc_name) # Left
        
        canvas.setFont('Helvetica', 6)
        canvas.drawRightString(7.5*inch, 11*inch, datetime.now().strftime('%Y-%m-%d')) # Right
        
        # Footer
        canvas.setFont('Helvetica', 9)
        canvas.drawCentredString(A4[0]/2, 0.75*inch, f"Page {doc.page}")
        
        canvas.restoreState()

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)

def create_epub(title, content, path):
    book = epub.EpubBook()
    book.set_identifier(title)
    book.set_title(title)
    book.set_language('es')

    c1 = epub.EpubHtml(title='Content', file_name='content.xhtml', lang='es')
    html_content = markdown2.markdown(content)
    c1.content = f"<h1>{title}</h1>{html_content}"
    
    book.add_item(c1)
    book.toc = (epub.Link('content.xhtml', 'Content', 'intro'),)
    book.spine = ['nav', c1]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    epub.write_epub(path, book)
